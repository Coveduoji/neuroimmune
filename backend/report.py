"""报告导出——按筛选条件生成专业安全报告，支持 docx / md / html 三种格式。

对标市面 SOC / 态势感知报告：封面 → 摘要 → 概览 → 图表 → 案件分析 → 免疫记忆 → 附录。
图表用 matplotlib 生成（HTML 内联 SVG，docx 嵌入 PNG，md 内嵌 base64 PNG）。
"""
from __future__ import annotations

import base64
import html
import io
import json
import sys
from collections import Counter
from datetime import datetime
from pathlib import Path

import matplotlib
matplotlib.use("Agg")  # 无 GUI 后端
import matplotlib.pyplot as plt

PROTO = str(Path(__file__).resolve().parent.parent / "prototype")
if PROTO not in sys.path:
    sys.path.insert(0, PROTO)

import db
import innate as innate_mod
import state
import tolerance as tol_mod

# 中文字体（按常见 CJK 字体优先级，找不到就回退 DejaVu）
plt.rcParams["font.sans-serif"] = [
    "Noto Sans CJK SC", "WenQuanYi Micro Hei", "SimHei",
    "Microsoft YaHei", "PingFang SC", "DejaVu Sans",
]
plt.rcParams["axes.unicode_minus"] = False

ACCENT = "#2a78d6"
TEAL = "#1baf7a"
CRITICAL = "#d03b3b"
WARNING = "#b7791f"
MUTED = "#8a8f98"

SUMMARY_SYSTEM = (
    "你是资深 SOC 安全分析师。根据给你的「统计概览 + 顶出案件摘要」，"
    "写一段 150~250 字的中文执行摘要（Executive Summary）：\n"
    "先一句话定性本次时间窗的安全态势，再说最值得关注的一个/几个案件，"
    "最后给一句处置建议。用事实、不编造，只输出摘要正文，不要标题、不要其他文字。"
)


def _normalize_datetime(s: str | None) -> str | None:
    """把前端 datetime-local 的 "YYYY-MM-DDTHH:MM" 归一化成 SQLite 的 "YYYY-MM-DD HH:MM:SS"。"""
    if not s:
        return None
    s = s.strip().replace("T", " ")
    if len(s) == 16:  # "YYYY-MM-DD HH:MM"
        s += ":00"
    return s


def _fig_bytes(fig, fmt: str) -> bytes:
    buf = io.BytesIO()
    fig.savefig(buf, format=fmt, bbox_inches="tight", dpi=150 if fmt == "png" else None)
    plt.close(fig)
    return buf.getvalue()


def build_report(filters: dict | None = None) -> dict:
    filters = filters or {}
    start = _normalize_datetime(filters.get("start"))
    end = _normalize_datetime(filters.get("end"))
    source = filters.get("source") or None

    alerts = db.list_alerts_report(start, end, source)
    surfaced = [a for a in alerts if not a["suppressed"]]
    suppressed_n = len(alerts) - len(surfaced)

    # 从告警反查案件（去重、保持强度排序）
    case_ids = sorted({a["case_id"] for a in surfaced if a["case_id"]})
    cases = []
    for cid in case_ids:
        c = db.get_case(cid)
        if not c:
            continue
        rpt = db.get_case_report(cid)
        cases.append({
            "correlation_uid": c["correlation_uid"], "title": c["title"],
            "strength": c["strength"], "status": c.get("status") or "New",
            "verdict": c.get("verdict") or "", "alerts": len(db.get_case_alerts(cid)),
            "entities": c.get("entities", []), "report": rpt,
        })
    # 按强度降序
    cases.sort(key=lambda x: x["strength"], reverse=True)

    escalated = [c for c in cases if c["report"]]
    verdict_filter = filters.get("verdict") or None
    status_filter = filters.get("status") or None
    if verdict_filter:
        cases = [c for c in cases if c["verdict"] == verdict_filter]
        escalated = [c for c in escalated if c["verdict"] == verdict_filter]
    if status_filter:
        cases = [c for c in cases if c["status"] == status_filter]
        escalated = [c for c in escalated if c["status"] == status_filter]

    source_counter = Counter(a["source"] for a in alerts)
    sources = [{"source": k, "count": v} for k, v in source_counter.most_common()]
    hour_counter = Counter((a.get("time") or "00:00")[:2] for a in alerts)
    hourly = [{"hour": h, "count": hour_counter.get(h, 0)} for h in [f"{i:02d}" for i in range(24)]]

    total = len(alerts)
    denoise = round((total - len(escalated)) / total * 100, 1) if total else 0.0

    report = {
        "title": "神经免疫安全报告",
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "time_range_label": _range_label(start, end),
        "filters": {k: v for k, v in filters.items() if v},
        "model_backend": _model_label(),
        "summary": "",
        "overview": {
            "alerts": total, "surfaced": len(surfaced), "suppressed": suppressed_n,
            "cases": len(cases), "escalated": len(escalated), "denoise": denoise,
            "tolerance": len(tol_mod.load_tolerance()),
            "innate": len(innate_mod.load_rules()),
        },
        "sources": sources,
        "hourly": hourly,
        "cases": cases,
        "escalated": escalated,
    }
    report["summary"] = _summary(report)
    return report


def _range_label(start, end) -> str:
    if not start and not end:
        return "全部时间"
    if start and end:
        return f"{start[:16]} ~ {end[:16]}"
    if start:
        return f"{start[:16]} 起"
    return f"截至 {end[:16]}"


def _model_label() -> str:
    mode = state.get_model_mode()
    if mode == "mock":
        return "Mock（规则版杏仁核）"
    m = state.get_model_config()
    return f"{m.get('model') or '真实模型'}（{mode}）"


def _stat_summary(o: dict, escalated: list) -> str:
    lines = [
        f"本时间窗共收到 {o['alerts']} 条告警，{o['suppressed']} 条被抑制，"
        f"{o['surfaced']} 条上板，聚合为 {o['cases']} 个案件，其中 {o['escalated']} 个顶出深析，"
        f"降噪率 {o['denoise']}%。",
    ]
    if escalated:
        top = escalated[0]
        v = (top["report"] or {}).get("verdict", "待定性")
        lines.append(
            f"最值得关注的是案件 {top['correlation_uid']}（强度 {top['strength']:.2f}，"
            f"{top['alerts']} 条告警），系统2 定性 {v}。"
        )
    else:
        lines.append("本时间窗没有顶出的高价值案件。")
    lines.append("建议对顶出案件逐条复核定性，并将误报回写免疫耐受、真阳性回写固有免疫，持续提升降噪与检测精度。")
    return "".join(lines)


def _summary(report: dict) -> str:
    o = report["overview"]
    # 有真实深想模型时生成专业摘要，失败/无模型则用统计摘要
    try:
        client = state.get_deep_client()
        if type(client).__name__ == "MockClient":
            raise ValueError("mock 模式，跳过 LLM 摘要")
        ctx = {
            "概览": o,
            "顶出案件": [
                {
                    "案件": c["correlation_uid"], "强度": c["strength"],
                    "定性": (c["report"] or {}).get("verdict", ""),
                    "摘要": ((c["report"] or {}).get("digest", "") or "")[:300],
                }
                for c in report["escalated"][:5]
            ],
        }
        raw = client.analyze(SUMMARY_SYSTEM + "\n\n" + json.dumps(ctx, ensure_ascii=False, indent=2))
        if raw.strip():
            return raw.strip()
    except Exception:
        pass
    return _stat_summary(o, report["escalated"])


# ---- 图表（matplotlib）----

def _chart_funnel(o: dict) -> plt.Figure:
    stages = ["告警", "上板", "案件", "顶出"]
    vals = [o["alerts"], o["surfaced"], o["cases"], o["escalated"]]
    fig, ax = plt.subplots(figsize=(6, 2.6))
    colors = [ACCENT, ACCENT, TEAL, CRITICAL]
    ax.barh(stages[::-1], vals[::-1], color=colors[::-1], height=0.55)
    for i, v in enumerate(vals):
        ax.text(v, 3 - i, f" {v}", va="center", fontsize=10)
    ax.set_xlabel("数量")
    ax.set_title("告警降噪漏斗")
    ax.spines[["top", "right"]].set_visible(False)
    return fig


def _chart_sources(sources: list) -> plt.Figure:
    top = sources[:12]
    fig, ax = plt.subplots(figsize=(6, max(2.2, 0.45 * len(top) + 1)))
    names = [s["source"] for s in reversed(top)]
    counts = [s["count"] for s in reversed(top)]
    ax.barh(names, counts, color=ACCENT, height=0.6)
    for i, v in enumerate(counts):
        ax.text(v, i, f" {v}", va="center", fontsize=9)
    ax.set_title("告警来源分布")
    ax.spines[["top", "right"]].set_visible(False)
    return fig


def _chart_hourly(hourly: list) -> plt.Figure:
    hours = [h["hour"] for h in hourly]
    counts = [h["count"] for h in hourly]
    fig, ax = plt.subplots(figsize=(7, 2.6))
    ax.bar(hours, counts, color=ACCENT, width=0.7)
    ax.set_xlabel("小时（当日）")
    ax.set_ylabel("告警数")
    ax.set_title("告警按小时分布")
    ax.spines[["top", "right"]].set_visible(False)
    return fig


def _chart_strength(cases: list) -> plt.Figure:
    strengths = [c["strength"] for c in cases] or [0]
    fig, ax = plt.subplots(figsize=(6, 2.6))
    ax.hist(strengths, bins=10, range=(0, 1.4), color=TEAL, edgecolor="white")
    ax.set_xlabel("案件强度")
    ax.set_ylabel("案件数")
    ax.set_title("案件强度分布")
    ax.spines[["top", "right"]].set_visible(False)
    return fig


def _charts_svg(report: dict) -> list[str]:
    out = []
    for f in (_chart_funnel(report["overview"]), _chart_sources(report["sources"]),
              _chart_hourly(report["hourly"]), _chart_strength(report["cases"])):
        out.append(_fig_bytes(f, "svg").decode("utf-8"))
    return out


def _charts_png(report: dict) -> list[bytes]:
    return [
        _fig_bytes(f, "png")
        for f in (_chart_funnel(report["overview"]), _chart_sources(report["sources"]),
                  _chart_hourly(report["hourly"]), _chart_strength(report["cases"]))
    ]


# ---- 渲染：HTML ----

_HTML_CSS = """
:root{--bg:#f4f5f7;--surface:#fff;--ink:#111;--ink2:#555;--muted:#888;--hairline:#e5e6e8;
--accent:#2a78d6;--teal:#1baf7a;--critical:#d03b3b;--warning:#b7791f}
*{box-sizing:border-box}body{margin:0;background:var(--bg);color:var(--ink);
font-family:system-ui,-apple-system,"Segoe UI","PingFang SC","Microsoft YaHei",sans-serif;line-height:1.6}
.wrap{max-width:960px;margin:0 auto;padding:32px 24px 64px}
h1{font-size:26px;margin:0 0 4px}h2{font-size:19px;margin:28px 0 10px;border-bottom:2px solid var(--hairline);padding-bottom:6px}
h3{font-size:15px;margin:16px 0 6px}
.card{background:var(--surface);border:1px solid var(--hairline);border-radius:10px;padding:18px 20px;margin-bottom:14px}
.kpis{display:grid;grid-template-columns:repeat(4,1fr);gap:10px}
.kpi{border:1px solid var(--hairline);border-radius:8px;padding:12px;text-align:center}
.kpi .v{font-size:24px;font-weight:700}.kpi .k{font-size:12px;color:var(--muted)}
.chart{margin:14px 0}.chart svg{max-width:100%;height:auto}
table{width:100%;border-collapse:collapse;font-size:13px}th,td{text-align:left;padding:8px 10px;border-bottom:1px solid var(--hairline)}
th{background:#fafafa;font-size:12px;color:var(--muted)}
.meta{color:var(--muted);font-size:12px}.tag{display:inline-block;padding:1px 8px;border-radius:99px;font-size:11px;font-weight:600}
.tag.tp{background:#fdeaea;color:var(--critical)}.tag.sus{background:#fdf3e0;color:var(--warning)}.tag.other{background:#eef0f3;color:var(--ink2)}
.chain{display:flex;align-items:center;gap:4px;overflow-x:auto;padding:6px 0}
.chain .node{background:#fafbfc;border:1px solid var(--hairline);border-radius:6px;padding:6px 10px;font-size:12px;white-space:nowrap}
.chain .arrow{color:var(--muted)}
"""


def _esc(s) -> str:
    return html.escape(str(s))


def _chain_html(attack_chain: list) -> str:
    if not attack_chain:
        return ""
    parts = []
    for i, a in enumerate(attack_chain):
        node = f'<div class="node"><b>{_esc(a.get("phase",""))}</b><br>{_esc(a.get("description",""))}</div>'
        parts.append(node)
        if i < len(attack_chain) - 1:
            parts.append('<div class="arrow">→</div>')
    return '<div class="chain">' + "".join(parts) + "</div>"


def render_html(report: dict) -> str:
    o = report["overview"]
    charts = _charts_svg(report)
    cases_html = []
    for c in report["escalated"]:
        r = c["report"] or {}
        verdict = r.get("verdict", "") or "待定性"
        tag_cls = "tp" if "True" in verdict or "Positive" in verdict else ("sus" if "Suspicious" in verdict else "other")
        cases_html.append(
            f'<h3>{_esc(c["correlation_uid"])} <span class="meta">强度 {c["strength"]:.2f} · {c["alerts"]} 条告警</span> '
            f'<span class="tag {tag_cls}">{_esc(verdict)}</span></h3>'
            f'<p>{_esc(r.get("digest", "") or "（无 AI 摘要）")}</p>'
            f'{_chain_html(r.get("attack_chain", []))}'
            f'{"<p><b>IOC：</b>" + "、".join(_esc(i["value"]) for i in r.get("iocs", [])) + "</p>" if r.get("iocs") else ""}'
            f'{"<p><b>处置建议：</b>" + "；".join(_esc(x) for x in r.get("remediations", [])) + "</p>" if r.get("remediations") else ""}'
        )
    case_rows = "".join(
        f'<tr><td><code>{_esc(c["correlation_uid"][:8])}</code></td><td>{c["strength"]:.2f}</td>'
        f'<td>{c["alerts"]}</td><td>{_esc(c["status"])}</td><td>{_esc(c["verdict"] or "—")}</td></tr>'
        for c in report["cases"]
    )
    chart_html = "".join(f'<div class="chart">{svg}</div>' for svg in charts)
    return (
        f'<!DOCTYPE html><html lang="zh-CN"><head><meta charset="utf-8">'
        f'<meta name="viewport" content="width=device-width,initial-scale=1">'
        f'<title>{_esc(report["title"])}</title><style>{_HTML_CSS}</style></head><body>'
        f'<div class="wrap">'
        f'<h1>{_esc(report["title"])}</h1>'
        f'<div class="meta">时间范围：{_esc(report["time_range_label"])} · 生成于 {_esc(report["generated_at"])} · 模型：{_esc(report["model_backend"])}</div>'
        f'<h2>摘要</h2><div class="card"><p>{_esc(report["summary"])}</p></div>'
        f'<h2>概览</h2><div class="kpis">'
        f'<div class="kpi"><div class="v">{o["alerts"]}</div><div class="k">告警</div></div>'
        f'<div class="kpi"><div class="v">{o["surfaced"]}</div><div class="k">上板</div></div>'
        f'<div class="kpi"><div class="v">{o["cases"]}</div><div class="k">案件</div></div>'
        f'<div class="kpi"><div class="v">{o["escalated"]}</div><div class="k">顶出（攻击链）</div></div>'
        f'</div>'
        f'<div class="card"><div class="meta">降噪率 <b>{o["denoise"]}%</b> · 被抑制 {o["suppressed"]} · 免疫耐受 {o["tolerance"]} · 固有免疫 {o["innate"]}</div></div>'
        f'<h2>图表</h2><div class="card">{chart_html}</div>'
        f'<h2>案件分析（{len(report["escalated"])} 个顶出案件）</h2>'
        f'<div class="card">{"".join(cases_html) or "<p class=meta>无顶出案件。</p>"}</div>'
        f'<h2>全部案件（{len(report["cases"])}）</h2><div class="card"><table>'
        f'<thead><tr><th>案件</th><th>强度</th><th>告警数</th><th>状态</th><th>定性</th></tr></thead><tbody>{case_rows}</tbody></table></div>'
        f'<h2>附录</h2><div class="card meta">筛选条件：{_esc(json.dumps(report["filters"], ensure_ascii=False) or "无")}<br>'
        f'本报告由神经免疫防御系统自动生成，图表为降噪漏斗、来源分布、小时分布、案件强度分布。</div>'
        f'</div></body></html>'
    )


# ---- 渲染：Markdown ----

def render_markdown(report: dict) -> str:
    o = report["overview"]
    lines = [
        f"# {report['title']}", "",
        f"> 时间范围：{report['time_range_label']} · 生成于 {report['generated_at']} · 模型：{report['model_backend']}", "",
        "## 摘要", "", report["summary"], "",
        "## 概览", "",
        f"- 告警 **{o['alerts']}**（上板 {o['surfaced']} · 抑制 {o['suppressed']}）",
        f"- 案件 **{o['cases']}** · 顶出（攻击链）**{o['escalated']}**",
        f"- 降噪率 **{o['denoise']}%** · 免疫耐受 {o['tolerance']} 条 · 固有免疫 {o['innate']} 条",
        "",
        "## 案件分析", "",
    ]
    for c in report["escalated"]:
        r = c["report"] or {}
        lines.append(f"### {c['correlation_uid']}（强度 {c['strength']:.2f} · {c['alerts']} 条告警 · {r.get('verdict','待定性')}）")
        if r.get("digest"):
            lines.append("")
            lines.append(r["digest"])
            lines.append("")
        for a in r.get("attack_chain", []):
            lines.append(f"- **{a.get('phase','')}**：{a.get('description','')}")
        if r.get("iocs"):
            lines.append(f"- IOC：{'、'.join(i['value'] for i in r['iocs'])}")
        if r.get("remediations"):
            lines.append(f"- 处置建议：{'；'.join(r['remediations'])}")
        lines.append("")
    lines += [
        "## 全部案件", "",
        "| 案件 | 强度 | 告警数 | 状态 | 定性 |",
        "|---|---|---|---|---|",
    ]
    for c in report["cases"]:
        lines.append(f"| `{c['correlation_uid'][:8]}` | {c['strength']:.2f} | {c['alerts']} | {c['status']} | {c['verdict'] or '—'} |")
    lines += ["", "---", "", "*本报告由神经免疫防御系统自动生成。*"]
    return "\n".join(lines)


# ---- 渲染：DOCX ----

def render_docx(report: dict) -> bytes:
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = docx.Document()
    o = report["overview"]

    title = doc.add_heading(report["title"], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"时间范围：{report['time_range_label']} · 生成于 {report['generated_at']} · 模型：{report['model_backend']}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_heading("摘要", level=1)
    doc.add_paragraph(report["summary"])

    doc.add_heading("概览", level=1)
    tbl = doc.add_table(rows=1, cols=4)
    hdr = tbl.rows[0].cells
    for i, txt in enumerate([f"告警 {o['alerts']}", f"上板 {o['surfaced']}", f"案件 {o['cases']}", f"顶出 {o['escalated']}"]):
        hdr[i].text = txt
    doc.add_paragraph(
        f"降噪率 {o['denoise']}% · 被抑制 {o['suppressed']} · 免疫耐受 {o['tolerance']} 条 · 固有免疫 {o['innate']} 条"
    )

    doc.add_heading("图表", level=1)
    for png in _charts_png(report):
        doc.add_picture(io.BytesIO(png), width=Inches(6.0))
        doc.add_paragraph()

    doc.add_heading(f"案件分析（{len(report['escalated'])} 个顶出案件）", level=1)
    if not report["escalated"]:
        doc.add_paragraph("无顶出案件。")
    for c in report["escalated"]:
        r = c["report"] or {}
        doc.add_heading(f"{c['correlation_uid']} · 强度 {c['strength']:.2f} · {c['alerts']} 条告警", level=2)
        doc.add_paragraph(f"定性：{r.get('verdict','待定性')}（置信度 {r.get('confidence','—')}）")
        if r.get("digest"):
            doc.add_paragraph(r["digest"])
        if r.get("attack_chain"):
            doc.add_paragraph("攻击链：")
            for a in r["attack_chain"]:
                doc.add_paragraph(f"• {a.get('phase','')}：{a.get('description','')}", style="List Bullet")
        if r.get("iocs"):
            doc.add_paragraph("IOC：" + "、".join(i["value"] for i in r["iocs"]))
        if r.get("remediations"):
            doc.add_paragraph("处置建议：" + "；".join(r["remediations"]))

    doc.add_heading("附录", level=1)
    doc.add_paragraph(f"筛选条件：{json.dumps(report['filters'], ensure_ascii=False) or '无'}")
    doc.add_paragraph("本报告由神经免疫防御系统自动生成。")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---- 入口 ----

FORMATS = {
    "html": ("text/html", "html", render_html),
    "md": ("text/markdown", "md", render_markdown),
    "docx": ("application/vnd.openxmlformats-officedocument.wordprocessingml.document", "docx", render_docx),
}


def export_report(filters: dict, fmt: str) -> tuple[str, str, bytes]:
    """返回 (media_type, ext, content)。"""
    if fmt not in FORMATS:
        raise ValueError(f"不支持的格式 {fmt!r}（支持 docx/md/html）")
    media, ext, renderer = FORMATS[fmt]
    report = build_report(filters)
    return media, ext, renderer(report)
